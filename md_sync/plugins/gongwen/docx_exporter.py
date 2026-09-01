"""Gongwen DOCX exporter — 生成符合 GB/T 9704-2012 的标准公文 Word 文档。

基础 docx 导出走 ``HTML → pandoc``，输出只是普通 Word 样式（无红头、无页码、
字体行距均非公文规格）。本导出器改用 python-docx **直接从文档结构构建** docx，
完整复现国标版式：

* A4，版心 156×225mm：天头 37 / 下 35 / 订口 28 / 翻口 26mm；
* 红头：发文机关标志（红色小标宋二号居中）+ 发文字号（三号仿宋居中）
  + 与版心等宽的红色分隔线；
* 标题：二号小标宋，居中，红色分隔线下空二行；
* 正文：三号仿宋_GB2312，首行缩进 2 字符，两端对齐，固定行距 28.8 磅
  （每面 22 行撑满版心），数字/字母 Times New Roman；
* 分级标题：一级黑体 / 二级楷体 / 三级、四级仿宋（三号）；
* 主送机关顶格、附件左空二字、落款署名右空二字、成文日期右空四字、附注左空二字；
* 版记：四号仿宋，左右各空一字，上下加分隔线；
* 页码：四号宋体 ``— n —``，一字线上距版心下边缘 7mm，单页码居右空一字、
  双页码居左空一字（奇偶页脚 + PAGE 域）。

python-docx 不可用时返回 False，由 pipeline 回退到基础 pandoc 导出
（与 PDF 导出器在无 PyMuPDF 时降级的模式一致）。
"""

from __future__ import annotations

import logging
import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor
from md_sync.core.md_engine import _MD
from md_sync.plugin.interface import DocxExporter

logger = logging.getLogger(__name__)

# ── GB/T 9704-2012 版式常量 ──────────────────────────────────────────────
_TOP_MM, _BOTTOM_MM, _LEFT_MM, _RIGHT_MM = 37.0, 35.0, 28.0, 26.0
_BODY_FONT_PT = 16.0  # 三号
_TITLE_FONT_PT = 22.0  # 二号
_PAGE_NO_FONT_PT = 14.0  # 四号
_LINE_EXACT_PT = 28.8  # 固定行距（每面 22 行撑满 225mm 版心）
_BODY_INDENT_PT = 2 * _BODY_FONT_PT  # 首行缩进 2 字符 = 32pt
_RED = "E60012"

# 字体：正文仿宋、一级黑体、二级楷体、标题/版头小标宋、页码宋体
_FS = "仿宋_GB2312"
_HT = "黑体"
_KT = "楷体_GB2312"
_XBS = "方正小标宋简体"
_ST = "宋体"
_LATIN = "Times New Roman"

_P_BDR_SUCCESSORS = (
    "w:shd",
    "w:tabs",
    "w:suppressAutoHyphens",
    "w:kinsoku",
    "w:wordWrap",
    "w:overflowPunct",
    "w:topLinePunct",
    "w:autoSpaceDE",
    "w:autoSpaceDN",
    "w:bidi",
    "w:adjustRightInd",
    "w:snapToGrid",
    "w:spacing",
    "w:ind",
    "w:contextualSpacing",
    "w:mirrorIndents",
    "w:suppressOverlap",
    "w:jc",
    "w:textDirection",
    "w:textAlignment",
    "w:textboxTightWrap",
    "w:outlineLvl",
    "w:divId",
    "w:cnfStyle",
    "w:rPr",
    "w:sectPr",
    "w:pPrChange",
)

_SIGN_DATE_RE = re.compile(r"^(?:\d{4}年\d{1,2}月\d{1,2}日|（这里是成文日期[^）]*）)$")
_ZHU_SONG_RE = re.compile(r"[：:]\s*$")
_ATTACH_RE = re.compile(r"^附件[：:]")
_NOTE_FULL_PAREN_RE = re.compile(r"^（[^（）]+）\s*$")


# ── 底层格式化工具 ───────────────────────────────────────────────────────


def _font(
    run,
    east_asia: str,
    size_pt: float,
    *,
    bold: bool = False,
    color: str | None = None,
    latin: str = _LATIN,
) -> None:
    """Set a run's font: Latin (Times New Roman) + CJK (eastAsia) + size/color."""
    run.font.name = latin
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:eastAsia"), east_asia)


def _fmt(
    p,
    *,
    align=None,
    line: float | None = _LINE_EXACT_PT,
    first_chars: int = 0,
    first_pt: float = 0.0,
    left_chars: int = 0,
    left_pt: float = 0.0,
    right_chars: int = 0,
    right_pt: float = 0.0,
    space_before: float = 0.0,
    space_after: float = 0.0,
    keep_with_next: bool = False,
    borders: list[tuple[str, int, str]] | None = None,
) -> None:
    """Apply GB/T 9704-2012 paragraph formatting (char-based indents, fixed line)."""
    pf = p.paragraph_format
    if align is not None:
        pf.alignment = align
    if line is not None:
        pf.line_spacing = Pt(line)
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    if space_before:
        pf.space_before = Pt(space_before)
    if space_after:
        pf.space_after = Pt(space_after)
    if keep_with_next:
        pf.keep_with_next = True

    p_pr = p._p.get_or_add_pPr()
    if first_chars or first_pt:
        ind = p_pr.get_or_add_ind()
        ind.set(qn("w:firstLineChars"), str(int(first_chars * 100)))
        ind.set(qn("w:firstLine"), str(int(first_pt * 20)))
    if left_chars or left_pt:
        ind = p_pr.get_or_add_ind()
        ind.set(qn("w:leftChars"), str(int(left_chars * 100)))
        ind.set(qn("w:left"), str(int(left_pt * 20)))
    if right_chars or right_pt:
        ind = p_pr.get_or_add_ind()
        ind.set(qn("w:rightChars"), str(int(right_chars * 100)))
        ind.set(qn("w:right"), str(int(right_pt * 20)))
    if borders:
        p_bdr = OxmlElement("w:pBdr")
        for side, sz, color in borders:
            b = OxmlElement(f"w:{side}")
            b.set(qn("w:val"), "single")
            b.set(qn("w:sz"), str(sz))  # 八分之一磅
            b.set(qn("w:space"), "1")
            b.set(qn("w:color"), color)
            p_bdr.append(b)
        p_pr.insert_element_before(p_bdr, *_P_BDR_SUCCESSORS)


# ── 块级解析（复用 markdown-it token 流，与 HTML raw 布局同源）──────────────


def _find_matching(tokens, start: int, open_type: str, close_type: str) -> int:
    depth = 0
    for k in range(start, len(tokens)):
        if tokens[k].type == open_type:
            depth += 1
        elif tokens[k].type == close_type:
            depth -= 1
            if depth == 0:
                return k
    return len(tokens) - 1


def _collect_bullets(tokens, start: int, end: int) -> list[str]:
    items: list[str] = []
    k = start
    while k < end:
        if tokens[k].type == "list_item_open":
            parts: list[str] = []
            m = k + 1
            depth = 0
            while m < end and not (tokens[m].type == "list_item_close" and depth == 0):
                tk = tokens[m].type
                if tk in ("bullet_list_open", "ordered_list_open"):
                    depth += 1
                elif tk in ("bullet_list_close", "ordered_list_close"):
                    depth -= 1
                elif tk == "inline":
                    parts.append(tokens[m].content.strip())
                m += 1
            text = " ".join(parts).strip()
            if text:
                items.append(text)
            k = m
        else:
            k += 1
    return items


def _table_rows(tokens, start: int, end: int) -> list[list[str]]:
    rows: list[list[str]] = []
    cur: list[str] | None = None
    for k in range(start, end):
        tk = tokens[k].type
        if tk == "tr_open":
            cur = []
        elif tk == "tr_close":
            if cur is not None:
                rows.append(cur)
            cur = None
        elif tk == "inline" and cur is not None:
            cur.append(tokens[k].content.strip())
    return rows


def _iter_blocks(source_raw: str):
    """Yield ``(kind, *data)`` tuples from raw Markdown (公文源文件)."""
    tokens = _MD.parse(source_raw)
    i, n = 0, len(tokens)
    while i < n:
        t = tokens[i]
        tt = t.type
        if tt == "heading_open":
            inline = tokens[i + 1] if i + 1 < n and tokens[i + 1].type == "inline" else None
            level = int(t.tag[1:])
            text = inline.content.strip() if inline else ""
            yield ("heading", level, text)
            i += 2 if inline else 1
        elif tt == "paragraph_open":
            inline = tokens[i + 1] if i + 1 < n and tokens[i + 1].type == "inline" else None
            text = inline.content.strip() if inline else ""
            if text:
                yield ("para", text)
            i += 2 if inline else 1
        elif tt in ("bullet_list_open", "ordered_list_open"):
            j = _find_matching(tokens, i, tt, tt.replace("_open", "_close"))
            ordered = tt == "ordered_list_open"
            items = _collect_bullets(tokens, i + 1, j)
            yield ("list", ordered, items)
            i = j + 1
        elif tt == "table_open":
            j = _find_matching(tokens, i, "table_open", "table_close")
            rows = _table_rows(tokens, i, j)
            yield ("table", rows)
            i = j + 1
        elif tt == "fence":
            lang = (t.info or "").strip().split()[0] if t.info else ""
            yield ("code", lang, t.content.rstrip("\n"))
            i += 1
        elif tt == "hr":
            yield ("hr",)
            i += 1
        elif tt == "blockquote_open":
            j = _find_matching(tokens, i, "blockquote_open", "blockquote_close")
            parts = [
                tokens[k].content.strip() for k in range(i, j + 1) if tokens[k].type == "inline"
            ]
            content = " ".join(parts).strip()
            if content:
                yield ("para", content)
            i = j + 1
        else:
            i += 1


# ── 公文语义分类（对应 filters.gongwen_chrome）─────────────────────────────


def _classify(blocks: list[tuple]) -> list[str]:
    """Assign each block a 公文 role.

    Roles: ``org`` / ``no`` (版头), ``title``, ``zhusong``, ``heading``,
    ``body``, ``attach``, ``sign`` / ``sign-date`` (落款), ``note`` (附注),
    ``hr`` / ``banji-hr``, ``banji`` (版记).
    """
    n = len(blocks)
    roles: list[str] = ["body"] * n

    title_idx = next(
        (i for i, b in enumerate(blocks) if b[0] == "heading" and b[1] == 1),
        None,
    )
    # 版记区域：最后一个 hr 之后的段落
    last_hr = next(
        (i for i in range(n - 1, -1, -1) if blocks[i][0] == "hr"),
        None,
    )
    banji_start = (last_hr + 1) if last_hr is not None else n

    for i in range(banji_start, n):
        if blocks[i][0] == "para":
            roles[i] = "banji"
    if last_hr is not None:
        roles[last_hr] = "banji-hr" if banji_start < n else "hr"

    if title_idx is not None:
        roles[title_idx] = "title"
        # 标题之前的段落 → 版头（首段红头、末段发文字号 + 分隔线）
        pre = [i for i in range(title_idx) if blocks[i][0] == "para"]
        if pre:
            roles[pre[0]] = "org"
            if len(pre) > 1:
                roles[pre[-1]] = "no"
        # 标题后第一个以全角/半角冒号结尾的段落 → 主送机关（顶格）
        zhusong = next(
            (i for i in range(title_idx + 1, n) if blocks[i][0] == "para" and roles[i] != "banji"),
            None,
        )
        if zhusong is not None and _ZHU_SONG_RE.search(blocks[zhusong][1]):
            roles[zhusong] = "zhusong"

    # 落款：正文区（版记之前）最后一段成文日期 + 其上一段署名
    end_scan = last_hr if last_hr is not None else n
    date_idx = next(
        (
            i
            for i in range(end_scan - 1, -1, -1)
            if blocks[i][0] == "para" and _SIGN_DATE_RE.match(blocks[i][1])
        ),
        None,
    )
    if date_idx is not None:
        roles[date_idx] = "sign-date"
        if (
            date_idx - 1 >= 0
            and blocks[date_idx - 1][0] == "para"
            and roles[date_idx - 1] in ("body",)
        ):
            roles[date_idx - 1] = "sign"
        # 成文日期之后的圆括号段落 → 附注（左空二字）
        for j in range(date_idx + 1, end_scan):
            if (
                blocks[j][0] == "para"
                and roles[j] == "body"
                and _NOTE_FULL_PAREN_RE.match(blocks[j][1])
            ):
                roles[j] = "note"

    # 附件：正文下空一行、左空二字
    for i, b in enumerate(blocks):
        if b[0] == "para" and roles[i] == "body" and _ATTACH_RE.match(b[1]):
            roles[i] = "attach"

    return roles


# ── 内容渲染 ─────────────────────────────────────────────────────────────


def _add_body_para(doc, text: str, role: str, t_fn) -> None:
    p = doc.add_paragraph()
    if role == "org":
        _fmt(
            p,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            line=_TITLE_FONT_PT + 12,
            space_before=28.8,
            space_after=4,
        )
        _font(p.add_run(t_fn(text)), _XBS, _TITLE_FONT_PT, color=_RED)
    elif role == "no":
        _fmt(
            p, align=WD_ALIGN_PARAGRAPH.CENTER, line=_LINE_EXACT_PT, borders=[("bottom", 12, _RED)]
        )
        _font(p.add_run(t_fn(text)), _FS, _BODY_FONT_PT)
    elif role == "title":
        # 标题：红色分隔线下空二行（前一“发文字号”段后补一空行，加上本段
        # space_before 各一行，兼顾 Word 相加 / LibreOffice 取大的差异）
        _fmt(
            p,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            line=_TITLE_FONT_PT + 12,
            space_before=_LINE_EXACT_PT,
            space_after=_LINE_EXACT_PT,
            keep_with_next=True,
        )
        _font(p.add_run(t_fn(text)), _XBS, _TITLE_FONT_PT)
    elif role == "zhusong":
        _fmt(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line=_LINE_EXACT_PT, keep_with_next=True)
        _font(p.add_run(t_fn(text)), _FS, _BODY_FONT_PT)
    elif role == "attach":
        _fmt(
            p,
            align=WD_ALIGN_PARAGRAPH.JUSTIFY,
            line=_LINE_EXACT_PT,
            left_chars=2,
            left_pt=_BODY_INDENT_PT,
            space_before=28.8,
        )
        _font(p.add_run(t_fn(text)), _FS, _BODY_FONT_PT)
    elif role == "sign":
        _fmt(
            p,
            align=WD_ALIGN_PARAGRAPH.RIGHT,
            line=_LINE_EXACT_PT,
            right_chars=2,
            right_pt=2 * _BODY_FONT_PT,
            space_before=28.8,
        )
        _font(p.add_run(t_fn(text)), _FS, _BODY_FONT_PT)
    elif role == "sign-date":
        _fmt(
            p,
            align=WD_ALIGN_PARAGRAPH.RIGHT,
            line=_LINE_EXACT_PT,
            right_chars=4,
            right_pt=4 * _BODY_FONT_PT,
        )
        _font(p.add_run(t_fn(text)), _FS, _BODY_FONT_PT)
    elif role == "note":
        _fmt(
            p,
            align=WD_ALIGN_PARAGRAPH.JUSTIFY,
            line=_LINE_EXACT_PT,
            left_chars=2,
            left_pt=_BODY_INDENT_PT,
        )
        _font(p.add_run(t_fn(text)), _FS, _BODY_FONT_PT)
    elif role == "banji":
        borders = []
        _fmt(
            p,
            align=WD_ALIGN_PARAGRAPH.JUSTIFY,
            line=_LINE_EXACT_PT,
            left_chars=1,
            left_pt=_PAGE_NO_FONT_PT,
            right_chars=1,
            right_pt=_PAGE_NO_FONT_PT,
            borders=borders,
        )
        _font(p.add_run(t_fn(text)), _FS, _PAGE_NO_FONT_PT)
    else:  # body
        _fmt(
            p,
            align=WD_ALIGN_PARAGRAPH.JUSTIFY,
            line=_LINE_EXACT_PT,
            first_chars=2,
            first_pt=_BODY_INDENT_PT,
        )
        _font(p.add_run(t_fn(text)), _FS, _BODY_FONT_PT)


def _add_heading(doc, text: str, level: int, t_fn) -> None:
    p = doc.add_paragraph()
    if level <= 2:
        # 一级标题：黑体，顶格
        _fmt(
            p,
            align=WD_ALIGN_PARAGRAPH.JUSTIFY,
            line=_LINE_EXACT_PT,
            space_before=14,
            keep_with_next=True,
        )
        _font(p.add_run(t_fn(text)), _HT, _BODY_FONT_PT)
    elif level == 3:
        # 二级标题：楷体，缩进 2 字符
        _fmt(
            p,
            align=WD_ALIGN_PARAGRAPH.JUSTIFY,
            line=_LINE_EXACT_PT,
            first_chars=2,
            first_pt=_BODY_INDENT_PT,
            space_before=14,
            keep_with_next=True,
        )
        _font(p.add_run(t_fn(text)), _KT, _BODY_FONT_PT)
    else:
        # 三级及以下：仿宋，缩进 2 字符
        _fmt(
            p,
            align=WD_ALIGN_PARAGRAPH.JUSTIFY,
            line=_LINE_EXACT_PT,
            first_chars=2,
            first_pt=_BODY_INDENT_PT,
            space_before=14,
            keep_with_next=True,
        )
        _font(p.add_run(t_fn(text)), _FS, _BODY_FONT_PT)


def _add_list(doc, ordered: bool, items: list[str], t_fn) -> None:
    for idx, item in enumerate(items, start=1):
        prefix = f"{idx}. " if ordered else "• "
        p = doc.add_paragraph()
        _fmt(
            p,
            align=WD_ALIGN_PARAGRAPH.JUSTIFY,
            line=_LINE_EXACT_PT,
            first_chars=2,
            first_pt=_BODY_INDENT_PT,
        )
        _font(p.add_run(prefix + t_fn(item)), _FS, _BODY_FONT_PT)


def _add_table(doc, rows: list[list[str]], t_fn) -> None:
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = doc.styles["Table Grid"]
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    col_w = Mm(156.0 / max(ncols, 1))
    for ci in range(ncols):
        for cell in table.columns[ci].cells:
            cell.width = col_w
    for ri, row in enumerate(rows):
        for ci in range(ncols):
            cell = table.cell(ri, ci)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            para = cell.paragraphs[0]
            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.line_spacing = 1.0
            run = para.add_run(t_fn(row[ci] if ci < len(row) else ""))
            _font(run, _FS, 12, bold=(ri == 0))
    # 表格后补一个空行，与正文分隔
    _fmt(doc.add_paragraph(), line=_LINE_EXACT_PT)


def _add_code(doc, text: str) -> None:
    p = doc.add_paragraph()
    _fmt(
        p,
        align=WD_ALIGN_PARAGRAPH.LEFT,
        line=None,
        space_before=7,
        space_after=7,
        borders=[
            ("top", 8, "D5D5D5"),
            ("left", 8, "D5D5D5"),
            ("bottom", 8, "D5D5D5"),
            ("right", 8, "D5D5D5"),
        ],
    )
    _font(p.add_run(text), "Courier New", 12, latin="Courier New")


def _add_hr(doc, role: str) -> None:
    p = doc.add_paragraph()
    if role == "banji-hr":
        # 版记顶部分隔线由首条版记行的上边框承担，这里省略
        return
    _fmt(p, align=WD_ALIGN_PARAGRAPH.LEFT, line=14, borders=[("top", 8, "000000")])


# ── 页码（奇偶页脚 + PAGE 域）─────────────────────────────────────────────


def _add_page_field(p) -> None:
    # OOXML 域结构：begin → instrText(PAGE) → separate → 缓存值 → end
    r = p.add_run()
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "begin")
    r._element.append(fld)
    _font(r, _ST, _PAGE_NO_FONT_PT)

    r = p.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    r._element.append(instr)
    _font(r, _ST, _PAGE_NO_FONT_PT)

    r = p.add_run()
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "separate")
    r._element.append(fld)
    _font(r, _ST, _PAGE_NO_FONT_PT)

    r = p.add_run("1")
    _font(r, _ST, _PAGE_NO_FONT_PT)

    r = p.add_run()
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "end")
    r._element.append(fld)
    _font(r, _ST, _PAGE_NO_FONT_PT)


def _fill_footer(footer, side: str) -> None:
    """单页码居右空一字 / 双页码居左空一字（四号宋体 ``— n —``）。"""
    p = footer.paragraphs[0]
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    _fmt(
        p,
        align=(WD_ALIGN_PARAGRAPH.RIGHT if side == "right" else WD_ALIGN_PARAGRAPH.LEFT),
        line=None,
        left_chars=(1 if side == "left" else 0),
        left_pt=(_PAGE_NO_FONT_PT if side == "left" else 0.0),
        right_chars=(1 if side == "right" else 0),
        right_pt=(_PAGE_NO_FONT_PT if side == "right" else 0.0),
    )
    _font(p.add_run("— "), _ST, _PAGE_NO_FONT_PT)
    _add_page_field(p)
    _font(p.add_run(" —"), _ST, _PAGE_NO_FONT_PT)


# ── 导出器 ───────────────────────────────────────────────────────────────


class GongwenDocxExporter(DocxExporter):
    """GB/T 9704-2012 公文 DOCX 导出器（红头 + 国标版式 + 一字线页码）。"""

    @property
    def name(self) -> str:
        return "gongwen"

    def export(
        self,
        doc,
        output_path,
        style_name: str = "",
        lang: str = "zh",
        translator=None,
    ) -> bool:
        source_raw = (doc.source_raw or "").strip()
        if not source_raw:
            logger.warning("[gongwen-docx] 无源文本（source_raw 为空），回退 pandoc")
            return False

        try:
            document = _build_document(doc, source_raw, lang, translator)
            output_path = Path(output_path)
            output_path.parent.mkdir(parents=True, exist_ok=True)
            document.save(str(output_path))
            logger.info("[gongwen-docx] ✓ 生成 %s", output_path.name)
            return True
        except Exception as e:
            logger.warning("[gongwen-docx] 生成失败，回退 pandoc: %s", e)
            return False


def _build_document(doc, source_raw: str, lang: str, translator):
    document = Document()
    _set_normal_style(document)

    # 版式：A4，版心 156×225mm，页脚距离校准至一字线上距版心下边缘 7mm
    sec = document.sections[0]
    sec.page_width = Mm(210)
    sec.page_height = Mm(297)
    sec.top_margin = Mm(_TOP_MM)
    sec.bottom_margin = Mm(_BOTTOM_MM)
    sec.left_margin = Mm(_LEFT_MM)
    sec.right_margin = Mm(_RIGHT_MM)
    sec.footer_distance = Mm(22.0)
    document.settings.odd_and_even_pages_header_footer = True
    _fill_footer(sec.footer, "right")
    _fill_footer(sec.even_page_footer, "left")

    def t_fn(text: str) -> str:
        if translator is None or lang == doc.source_lang:
            return text
        cached = translator.lookup(text, lang)
        return cached if cached else text

    blocks = list(_iter_blocks(source_raw))
    roles = _classify(blocks)

    # 版记行：首条加粗顶线、末条加粗底线、中间用细线分隔（GB/T 9704-2012 §7.4）
    banji_idx = [i for i, r in enumerate(roles) if r == "banji"]

    for i, (block, role) in enumerate(zip(blocks, roles, strict=True)):
        kind = block[0]
        if kind == "para":
            if role == "banji":
                p = document.add_paragraph()
                first = (i == banji_idx[0]) if banji_idx else False
                last = (i == banji_idx[-1]) if banji_idx else False
                borders = []
                if first:
                    borders.append(("top", 8, "000000"))
                # 末条用粗线（0.35mm≈sz8），中间用细线（0.25mm≈sz6）
                borders.append(("bottom", 8 if last else 6, "000000"))
                _fmt(
                    p,
                    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                    line=_LINE_EXACT_PT,
                    left_chars=1,
                    left_pt=_PAGE_NO_FONT_PT,
                    right_chars=1,
                    right_pt=_PAGE_NO_FONT_PT,
                    borders=borders,
                )
                _font(p.add_run(t_fn(block[1])), _FS, _PAGE_NO_FONT_PT)
            else:
                _add_body_para(document, block[1], role, t_fn)
                # 红色分隔线下空二行：发文字号段之后补一行空行
                if role == "no":
                    _fmt(document.add_paragraph(), line=_LINE_EXACT_PT)
        elif kind == "heading":
            if role == "title":
                _add_body_para(document, block[2], "title", t_fn)
            else:
                _add_heading(document, block[2], block[1], t_fn)
        elif kind == "list":
            _add_list(document, block[1], block[2], t_fn)
        elif kind == "table":
            _add_table(document, block[1], t_fn)
        elif kind == "code":
            _add_code(document, block[2])
        elif kind == "hr":
            _add_hr(document, role)

    return document


def _set_normal_style(document) -> None:
    """Normal 样式：三号仿宋、固定行距 28.8 磅、两端对齐。"""
    normal = document.styles["Normal"]
    normal.font.name = _LATIN
    normal.font.size = Pt(_BODY_FONT_PT)
    r_fonts = normal._element.get_or_add_rPr().get_or_add_rFonts()
    r_fonts.set(qn("w:eastAsia"), _FS)
    pf = normal.paragraph_format
    pf.line_spacing = Pt(_LINE_EXACT_PT)
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.space_after = Pt(0)
